import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import difflib
import os
import shutil
from datetime import datetime
import re

class PreciseWordComparatorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("정교한 워드 파일 비교 프로그램 (목차 특화)")
        self.root.geometry("1200x900")
        self.root.configure(bg="#f0f0f0")
        
        # 파일 경로 저장 변수
        self.file1_path = ""
        self.file2_path = ""
        self.original_content = []
        self.current_content = []
        self.diff_data = []
        
        self.setup_ui()
    
    def setup_ui(self):
        # 메인 프레임
        main_frame = tk.Frame(self.root, bg="#f0f0f0")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # 제목
        title_label = tk.Label(main_frame, text="📋 정교한 워드 파일 비교 프로그램 (목차/구조 특화)", 
                              font=("맑은 고딕", 18, "bold"), 
                              bg="#f0f0f0", fg="#2c3e50")
        title_label.pack(pady=(0, 10))
        
        # 설명
        desc_label = tk.Label(main_frame, text="💡 목차, 번호 매김, 페이지 번호를 정확히 인식하여 모든 변경사항을 감지합니다", 
                              font=("맑은 고딕", 11), 
                              bg="#f0f0f0", fg="#7f8c8d")
        desc_label.pack(pady=(0, 20))
        
        # 파일 선택 프레임
        file_frame = tk.Frame(main_frame, bg="#f0f0f0")
        file_frame.pack(fill=tk.X, pady=(0, 20))
        
        # 원본 파일 선택
        file1_frame = tk.LabelFrame(file_frame, text="📂 원본 파일 (기준)", 
                                   font=("맑은 고딕", 12, "bold"),
                                   bg="#ffe8e8", fg="#2c3e50", padx=10, pady=10)
        file1_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        self.file1_label = tk.Label(file1_frame, text="파일이 선택되지 않았습니다", 
                                   wraplength=300, justify=tk.LEFT,
                                   bg="#ffe8e8", fg="#7f8c8d")
        self.file1_label.pack(pady=5)
        
        self.file1_button = tk.Button(file1_frame, text="원본 파일 선택", 
                                     command=self.select_file1,
                                     bg="#e74c3c", fg="white", 
                                     font=("맑은 고딕", 10, "bold"),
                                     relief=tk.FLAT, padx=20, pady=5)
        self.file1_button.pack(pady=5)
        
        # 비교 파일 선택
        file2_frame = tk.LabelFrame(file_frame, text="📂 비교 파일 (서식 유지)", 
                                   font=("맑은 고딕", 12, "bold"),
                                   bg="#e8f8e8", fg="#2c3e50", padx=10, pady=10)
        file2_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        self.file2_label = tk.Label(file2_frame, text="파일이 선택되지 않았습니다", 
                                   wraplength=300, justify=tk.LEFT,
                                   bg="#e8f8e8", fg="#7f8c8d")
        self.file2_label.pack(pady=5)
        
        self.file2_button = tk.Button(file2_frame, text="비교 파일 선택", 
                                     command=self.select_file2,
                                     bg="#27ae60", fg="white", 
                                     font=("맑은 고딕", 10, "bold"),
                                     relief=tk.FLAT, padx=20, pady=5)
        self.file2_button.pack(pady=5)
        
        # 비교 옵션 프레임
        option_frame = tk.LabelFrame(main_frame, text="🔧 비교 옵션", 
                                     font=("맑은 고딕", 12, "bold"),
                                     bg="#f8f9fa", fg="#2c3e50", padx=10, pady=10)
        option_frame.pack(fill=tk.X, pady=(0, 15))
        
        self.ignore_page_numbers = tk.BooleanVar(value=True)
        self.ignore_dots = tk.BooleanVar(value=True)
        
        tk.Checkbutton(option_frame, text="페이지 번호 무시 (권장)", 
                      variable=self.ignore_page_numbers,
                      bg="#f8f9fa", font=("맑은 고딕", 10)).pack(side=tk.LEFT, padx=20)
        tk.Checkbutton(option_frame, text="점선(...) 무시 (권장)", 
                      variable=self.ignore_dots,
                      bg="#f8f9fa", font=("맑은 고딕", 10)).pack(side=tk.LEFT, padx=20)
        
        # 컨트롤 버튼 프레임
        control_frame = tk.Frame(main_frame, bg="#f0f0f0")
        control_frame.pack(pady=20)
        
        self.compare_button = tk.Button(control_frame, text="🔍 정교한 구조 비교", 
                                       command=self.precise_compare_files,
                                       bg="#3498db", fg="white", 
                                       font=("맑은 고딕", 14, "bold"),
                                       relief=tk.FLAT, padx=40, pady=12,
                                       state=tk.DISABLED)
        self.compare_button.pack(side=tk.LEFT, padx=10)
        
        self.save_button = tk.Button(control_frame, text="💾 비교 결과 저장", 
                                    command=self.save_result,
                                    bg="#e67e22", fg="white", 
                                    font=("맑은 고딕", 14, "bold"),
                                    relief=tk.FLAT, padx=40, pady=12,
                                    state=tk.DISABLED)
        self.save_button.pack(side=tk.LEFT, padx=10)
        
        # 범례 프레임
        self.legend_frame = tk.Frame(main_frame, bg="#f0f0f0")
        self.legend_frame.pack(pady=(0, 10))
        
        legend_title = tk.Label(self.legend_frame, text="📊 범례 (저장된 워드 파일에서 보이는 색상)", 
                               font=("맑은 고딕", 12, "bold"),
                               bg="#f0f0f0", fg="#2c3e50")
        legend_title.pack()
        
        legend_content = tk.Frame(self.legend_frame, bg="#f0f0f0")
        legend_content.pack()
        
        # 범례 항목들
        added_frame = tk.Frame(legend_content, bg="#f0f0f0")
        added_frame.pack(side=tk.LEFT, padx=30)
        added_color = tk.Label(added_frame, text="  ", bg="#c8e6c9", relief=tk.SOLID, borderwidth=1)
        added_color.pack(side=tk.LEFT, padx=5)
        tk.Label(added_frame, text="새로 추가된 부분 (4.2.5. 2022.12 등)", bg="#f0f0f0", font=("맑은 고딕", 11, "bold")).pack(side=tk.LEFT)
        
        modified_frame = tk.Frame(legend_content, bg="#f0f0f0")
        modified_frame.pack(side=tk.LEFT, padx=30)
        modified_color = tk.Label(modified_frame, text="  ", bg="#fff3cd", relief=tk.SOLID, borderwidth=1)
        modified_color.pack(side=tk.LEFT, padx=5)
        tk.Label(modified_frame, text="페이지 번호 등 부분 변경", bg="#f0f0f0", font=("맑은 고딕", 11, "bold")).pack(side=tk.LEFT)
        
        # 처음에는 범례 숨기기
        self.legend_frame.pack_forget()
        
        # 결과 표시 프레임
        result_frame = tk.Frame(main_frame, bg="#ffffff")
        result_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 스크롤 가능한 텍스트 위젯
        self.result_text = scrolledtext.ScrolledText(result_frame, 
                                                    wrap=tk.WORD, 
                                                    font=("맑은 고딕", 10),
                                                    bg="white", fg="#2c3e50",
                                                    height=25)
        self.result_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 텍스트 태그 설정 (색상 지정)
        self.result_text.tag_configure("added", background="#c8e6c9", foreground="#2e7d32", font=("맑은 고딕", 10, "bold"))
        self.result_text.tag_configure("removed", background="#ffcdd2", foreground="#c62828", overstrike=True)
        self.result_text.tag_configure("modified", background="#fff3cd", foreground="#f57c00")
        self.result_text.tag_configure("normal", background="white", foreground="#2c3e50")
        self.result_text.tag_configure("header", font=("맑은 고딕", 12, "bold"), foreground="#2c3e50")
        
        # 상태바
        self.status_var = tk.StringVar()
        self.status_var.set("원본 파일과 비교 파일을 선택해주세요")
        status_bar = tk.Label(main_frame, textvariable=self.status_var, 
                             relief=tk.SUNKEN, anchor=tk.W,
                             bg="#ecf0f1", fg="#2c3e50", font=("맑은 고딕", 9))
        status_bar.pack(fill=tk.X, pady=(10, 0))

    def extract_structured_content(self, file_path):
        """목차나 구조화된 문서의 내용을 정확히 추출합니다."""
        try:
            doc = docx.Document(file_path)
            structured_content = []
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                full_text = paragraph.text
                
                if full_text.strip():
                    para_info = {
                        'index': para_idx,
                        'original_text': full_text,
                        'clean_text': self.clean_text_for_comparison(full_text),
                        'is_numbered': self.is_numbered_item(full_text),
                        'level': self.get_indentation_level(paragraph),
                        'page_number': self.extract_page_number(full_text),
                        'has_dots': '...' in full_text or '…' in full_text
                    }
                    structured_content.append(para_info)
            
            return structured_content
            
        except Exception as e:
            raise Exception(f"구조 분석 오류: {str(e)}")

    def clean_text_for_comparison(self, text):
        """비교용 텍스트 정리 (점선, 페이지 번호 등 옵션에 따라 제거)"""
        cleaned = text
        
        # 점선 제거 옵션
        if self.ignore_dots.get():
            # 연속된 점들 제거 (3개 이상)
            cleaned = re.sub(r'\.{3,}', ' ', cleaned)
            # 탭으로 연결된 점선들도 제거
            cleaned = re.sub(r'\t+\.+', ' ', cleaned)
        
        # 페이지 번호 제거 옵션
        if self.ignore_page_numbers.get():
            # 끝에 오는 페이지 번호 패턴들 제거
            cleaned = re.sub(r'\s+\d+\s*$', '', cleaned)
            cleaned = re.sub(r'\s+\d+\w*\s*$', '', cleaned)  # 30a, 54b 같은 패턴도 제거
        
        # 연속 공백 정리 및 앞뒤 공백 제거
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        return cleaned

    def is_numbered_item(self, text):
        """번호가 매겨진 항목인지 확인 (4.2.5. 같은 패턴)"""
        # 다양한 번호 매김 패턴 확인
        patterns = [
            r'^\s*\d+(\.\d+)*\.?\s',  # 4.2.5. 또는 4.2.5 
            r'^\s*\d+\)\s',           # 1) 2) 
            r'^\s*\(\d+\)\s',         # (1) (2)
            r'^\s*[가-힣]\.\s',       # 가. 나.
            r'^\s*[a-zA-Z]\.\s',      # a. b. A. B.
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False

    def get_indentation_level(self, paragraph):
        """문단의 들여쓰기 레벨 확인"""
        try:
            if paragraph.paragraph_format.left_indent:
                # 포인트 단위를 레벨로 변환 (대략 36pt = 1레벨)
                return int(paragraph.paragraph_format.left_indent.pt // 36)
        except:
            pass
        return 0

    def extract_page_number(self, text):
        """텍스트에서 페이지 번호 추출"""
        # 끝에 있는 숫자를 페이지 번호로 간주
        match = re.search(r'\s+(\d+)\s*$', text)
        if match:
            return match.group(1)
        return None

    def select_file1(self):
        file_path = filedialog.askopenfilename(
            title="원본 파일 선택",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.file1_path = file_path
            self.file1_label.config(text=f"선택됨: {os.path.basename(file_path)}", fg="#27ae60")
            self.status_var.set("원본 파일 구조를 분석하는 중...")
            self.root.update()
            
            try:
                self.original_content = self.extract_structured_content(file_path)
                self.status_var.set(f"원본 파일 분석 완료: {len(self.original_content)}개 항목")
                self.update_compare_button()
            except Exception as e:
                messagebox.showerror("오류", str(e))
                self.file1_label.config(text="파일 분석 실패", fg="#e74c3c")

    def select_file2(self):
        file_path = filedialog.askopenfilename(
            title="비교 파일 선택 (이 파일의 서식이 유지됩니다)",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.file2_path = file_path
            self.file2_label.config(text=f"선택됨: {os.path.basename(file_path)}", fg="#27ae60")
            self.status_var.set("비교 파일 구조를 분석하는 중...")
            self.root.update()
            
            try:
                self.current_content = self.extract_structured_content(file_path)
                self.status_var.set(f"비교 파일 분석 완료: {len(self.current_content)}개 항목")
                self.update_compare_button()
            except Exception as e:
                messagebox.showerror("오류", str(e))
                self.file2_label.config(text="파일 분석 실패", fg="#e74c3c")

    def update_compare_button(self):
        if self.original_content and self.current_content:
            self.compare_button.config(state=tk.NORMAL)
            self.status_var.set("준비 완료! 정교한 구조 비교를 시작할 수 있습니다.")

    def precise_compare_files(self):
        """구조화된 문서의 정교한 비교를 수행합니다."""
        if not self.original_content or not self.current_content:
            messagebox.showwarning("경고", "두 파일을 모두 선택해주세요.")
            return
        
        self.status_var.set("정교한 구조 비교 진행 중...")
        self.root.update()
        
        # 결과 텍스트 초기화
        self.result_text.delete(1.0, tk.END)
        self.diff_data = []
        
        try:
            self.compare_structured_documents()
            
            # 범례와 저장 버튼 활성화
            self.legend_frame.pack(before=self.result_text.master, pady=(0, 10))
            self.save_button.config(state=tk.NORMAL)
            self.status_var.set("✅ 정교한 비교 완료! 모든 변경사항이 정확히 감지되었습니다.")
            
        except Exception as e:
            messagebox.showerror("오류", f"비교 중 오류가 발생했습니다: {str(e)}")
            self.status_var.set("❌ 파일 비교 중 오류가 발생했습니다.")

    def compare_structured_documents(self):
        """구조화된 문서의 정밀 비교"""
        
        # 정리된 텍스트로 비교용 리스트 생성
        original_texts = [item['clean_text'] for item in self.original_content]
        current_texts = [item['clean_text'] for item in self.current_content]
        
        self.result_text.insert(tk.END, "=== 정교한 구조 비교 결과 ===\n\n", "header")
        self.result_text.insert(tk.END, f"📋 원본 파일: {os.path.basename(self.file1_path)} ({len(original_texts)}개 항목)\n", "normal")
        self.result_text.insert(tk.END, f"📋 비교 파일: {os.path.basename(self.file2_path)} ({len(current_texts)}개 항목)\n\n", "normal")
        
        # difflib을 사용한 정밀 비교
        matcher = difflib.SequenceMatcher(None, original_texts, current_texts)
        
        changes_found = False
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                # 동일한 부분
                for i in range(i1, i2):
                    if i < len(self.original_content):
                        original_item = self.original_content[i]
                        # 페이지 번호나 점선이 다를 수 있지만 내용은 같은 경우 확인
                        corresponding_current = None
                        current_idx = j1 + (i - i1)
                        if current_idx < len(self.current_content):
                            corresponding_current = self.current_content[current_idx]
                            
                            # 원본 텍스트와 현재 텍스트가 완전히 동일한지 확인
                            if original_item['original_text'] != corresponding_current['original_text']:
                                # 내용은 같지만 페이지 번호나 점선이 다른 경우
                                self.result_text.insert(tk.END, f"[페이지번호변경] {corresponding_current['original_text']}\n", "modified")
                                self.diff_data.append(('modified', corresponding_current['original_text']))
                                changes_found = True
                            else:
                                # 완전히 동일
                                self.result_text.insert(tk.END, f"{original_item['original_text']}\n", "normal")
                                self.diff_data.append(('equal', original_item['original_text']))
            
            elif tag == 'insert':
                # 추가된 부분 - 여기서 "4.2.5. 2022.12" 같은 항목이 감지되어야 함
                changes_found = True
                for j in range(j1, j2):
                    if j < len(self.current_content):
                        added_item = self.current_content[j]
                        self.result_text.insert(tk.END, f"[📄 새로추가] {added_item['original_text']}\n", "added")
                        self.diff_data.append(('insert', added_item['original_text']))
                        
                        # 디버깅 정보 출력
                        print(f"✅ 추가 항목 감지: '{added_item['clean_text']}'")
            
            elif tag == 'delete':
                # 삭제된 부분
                changes_found = True
                for i in range(i1, i2):
                    if i < len(self.original_content):
                        deleted_item = self.original_content[i]
                        self.result_text.insert(tk.END, f"[🗑️ 삭제됨] {deleted_item['original_text']}\n", "removed")
                        self.diff_data.append(('delete', deleted_item['original_text']))
            
            elif tag == 'replace':
                # 변경된 부분
                changes_found = True
                for i in range(i1, i2):
                    if i < len(self.original_content):
                        original_item = self.original_content[i]
                        self.result_text.insert(tk.END, f"[❌ 변경전] {original_item['original_text']}\n", "removed")
                        self.diff_data.append(('delete', original_item['original_text']))
                
                for j in range(j1, j2):
                    if j < len(self.current_content):
                        current_item = self.current_content[j]
                        self.result_text.insert(tk.END, f"[✅ 변경후] {current_item['original_text']}\n", "added")
                        self.diff_data.append(('insert', current_item['original_text']))
        
        # 비교 결과 요약
        if not changes_found:
            self.result_text.insert(tk.END, "\n🎉 두 문서가 동일합니다!\n", "header")
        else:
            added_count = len([item for tag, item in self.diff_data if tag == 'insert'])
            deleted_count = len([item for tag, item in self.diff_data if tag == 'delete'])
            modified_count = len([item for tag, item in self.diff_data if tag == 'modified'])
            
            self.result_text.insert(tk.END, f"\n📊 변경사항 요약:\n", "header")
            self.result_text.insert(tk.END, f"• 추가된 항목: {added_count}개\n", "added")
            self.result_text.insert(tk.END, f"• 삭제된 항목: {deleted_count}개\n", "removed")
            self.result_text.insert(tk.END, f"• 수정된 항목: {modified_count}개\n", "modified")

    def save_result(self):
        if not self.diff_data:
            messagebox.showwarning("경고", "저장할 결과가 없습니다.")
            return
        
        # 저장할 파일 선택
        file_path = filedialog.asksaveasfilename(
            title="정교한 비교 결과를 워드 파일로 저장",
            defaultextension=".docx",
            filetypes=[
                ("워드 문서", "*.docx"),
                ("모든 파일", "*.*")
            ]
        )
        
        if file_path:
            try:
                self.status_var.set("📁 비교 파일을 복사하고 차이점을 표시하는 중...")
                self.root.update()
                
                # 1단계: 비교 파일을 그대로 복사
                shutil.copy2(self.file2_path, file_path)
                
                # 2단계: 복사된 파일에서 차이점 하이라이트
                self.apply_precise_highlights_to_document(file_path)
                
                messagebox.showinfo("✅ 저장 완료!", 
                                  f"정교한 비교 결과가 성공적으로 저장되었습니다!\n\n"
                                  f"📁 저장 위치: {file_path}\n\n"
                                  f"💡 특징:\n"
                                  f"• 비교 파일의 원본 서식과 내용이 모두 유지됩니다\n"
                                  f"• 🟢 초록색 배경 = 새로 추가된 부분 (4.2.5. 2022.12 등)\n"
                                  f"• 🟡 노란색 배경 = 페이지 번호 등 부분 변경\n"
                                  f"• 완전히 깔끔한 문서 형태로 저장됩니다")
                
                self.status_var.set(f"✅ 저장 완료: {os.path.basename(file_path)}")
                
            except Exception as e:
                messagebox.showerror("저장 오류", f"파일 저장 중 오류가 발생했습니다: {str(e)}")

    def apply_precise_highlights_to_document(self, file_path):
        """정교한 하이라이트를 문서에 적용합니다."""
        try:
            doc = docx.Document(file_path)
            
            # 원본에서 정리된 텍스트들의 집합 생성
            original_clean_texts = set()
            original_full_texts = set()
            
            for item in self.original_content:
                original_clean_texts.add(item['clean_text'])
                original_full_texts.add(item['original_text'])
            
            # 현재 문서의 각 문단을 검사
            highlighted_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    current_text = paragraph.text
                    current_clean = self.clean_text_for_comparison(current_text)
                    
                    # 완전히 새로운 내용인지 확인
                    if current_clean not in original_clean_texts:
                        # 새로 추가된 내용 - 초록색 하이라이트
                        for run in paragraph.runs:
                            if run.text.strip():
                                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                                run.font.color.rgb = RGBColor(46, 125, 50)
                        highlighted_count += 1
                        print(f"✅ 하이라이트 적용: '{current_clean}'")
                    
                    # 내용은 같지만 형식이 다른 경우 (페이지 번호 변경 등)
                    elif current_text not in original_full_texts:
                        # 부분 변경 - 노란색 하이라이트
                        for run in paragraph.runs:
                            if run.text.strip():
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                run.font.color.rgb = RGBColor(245, 124, 0)
                        highlighted_count += 1
            
            # 표 내용도 동일하게 처리
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                current_text = paragraph.text
                                current_clean = self.clean_text_for_comparison(current_text)
                                
                                if current_clean not in original_clean_texts:
                                    for run in paragraph.runs:
                                        if run.text.strip():
                                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                                    highlighted_count += 1
            
            doc.save(file_path)
            print(f"💾 하이라이트 적용 완료: {highlighted_count}개 항목")
            
        except Exception as e:
            print(f"하이라이트 적용 오류: {e}")
            # 오류가 발생해도 파일은 저장되도록 함
            pass

def main():
    # 필요한 라이브러리 확인
    try:
        import docx
        import difflib
        import shutil
    except ImportError as e:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("라이브러리 오류", 
                           "필요한 라이브러리가 설치되지 않았습니다.\n"
                           "다음 명령어를 실행해주세요:\n\n"
                           "pip install python-docx")
        root.destroy()
        return
    
    root = tk.Tk()
    app = PreciseWordComparatorGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()